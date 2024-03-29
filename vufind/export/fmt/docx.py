#!/usr/bin/python3
#
# Vufind - Export DOCX format module
#

from io import BytesIO
from datetime import datetime

#DOCX

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# VAR

VUFIND='https://vufind.ucl.cas.cz/Record/'
WARN="Využitím zdrojů České literární bibliografie se uživatel zavazuje odkázat na její využití v každé publikaci, kvalifikační práci či jiném výstupu, a to následující formou: 'Při vzniku práce [knihy/studie/...] byly využity zdroje výzkumné infrastruktury Česká literární bibliografie – https://clb.ucl.cas.cz/ (kód ORJ: 90136).'"
ADDRESS='Česká literární bibliografie © ' + datetime.now().strftime('%Y') +  ' clb@ucl.cas.cz Na Florenci, 1420/3, 110 00 Praha'

RESOURCE = {
	'Literary Samizdat Bibliography':'Bibliografie samizdatu',
	'Literary Web Bibliography':'Bibliografie internetu',
	'Literary Exile Bibliography':'Bibliografie exilu',
	'Retrospective Bibliography (up to 1945)':'Retrospektivní bibliografie (do roku 1945)',
	'Current Bibliography (after 1945)':'Současná bibliografie (po roce 1945)',
	'ICL Catalogue':'Katalog UCL',
	'Database of Processed Journals':'Databáze excerpovaných časopisů',
	'Czech Literary Authorities Database':'České literární osobnosti',
	'Polonica in Czech Samizdat':'Polonika v českém samizdatu',
	'Database of Foreign Bohemica':'Databáze zahraničních bohemik',
	'Book Series Database':'Databáze knižních edic',
	'Czech Literature in Translation':'Databáze překladů české literatury'
}

# DEF

def prep(text):
	for prep in (' u ',' k ',' o ',' s ',' v ',' z ',' a ',' i '):
		if prep in text: text = text.replace(prep, prep.rstrip() + '\u00a0')
	return text

def name_to_upper(name):
	n = name.strip(', ')
	if len(n.split(',')) == 2:
		return n.split(',')[0].strip().upper() + ', ' + n.split(',')[1].strip() 
	return n

def name_to_upper(name):
	n = name.strip(', ')
	if len(n.split(',')) == 2:
		return n.split(',')[0].strip().upper() + ', ' + n.split(',')[1].strip() 
	return n

def hyperlink(par, text, url):
	part = par.part# document.xml.rels => relation ID
	rid = part.relate_to(url, RT.HYPERLINK, is_external=True)

	hyperlink = OxmlElement('w:hyperlink')
	hyperlink.set(qn('r:id'), rid, )

	run = OxmlElement('w:r')
	rpr = OxmlElement('w:rPr')

	run.append(rpr)
	run.text = text

	hyperlink.append(run)
	par._p.append(hyperlink)

	return hyperlink

def docx(data, lang):
	ret = BytesIO()
	# init
	doc = Document()
	sections = doc.sections
	for section in sections:
		section.left_margin = Cm(1.5)
		section.right_margin = Cm(1.5)
	# head
	par = doc.add_paragraph(WARN)
	par = doc.add_paragraph()
	#header = sections[0].header
	#par = header.paragraphs[0]
	#run = par.add_run()
	#run.add_picture(LOGO,width=Inches(1.0))
	#par.text = HEADER
	# data
	for record in data['response']['docs']:
		par = doc.add_paragraph()
		if 'info_resource_str_mv' in record:
			resource = record['info_resource_str_mv'][0]
			if lang == 'cs': resource = RESOURCE[record['info_resource_str_mv'][0]]
			par.add_run(resource + ', ')
			hyperlink(par, record['id'], VUFIND + record['id'])
		else:
			hyperlink(par, record['id'], VUFIND + record['id'])
		par.paragraph_format.keep_with_next = True
		par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		if 'export_100a_str' in record:
			par = doc.add_paragraph()
			par.add_run(name_to_upper(record['export_100a_str'])).bold = True
			if 'export_100bc_str' in record:
				par.add_run(name_to_upper(record['export_100bc_str']))
			par.paragraph_format.keep_with_next = True
		if 'export_245_str' in record:
			par = doc.add_paragraph(prep(record['export_245_str']))
			par.paragraph_format.keep_with_next = True
		if 'export_260264_str_mv' in record:
			par = doc.add_paragraph(prep(' '.join(record['export_260264_str_mv'])))
			par.paragraph_format.keep_with_next = True
		if 'export_490_str_mv' in record:
			par = doc.add_paragraph('(' + prep(' '.join(record['export_490_str_mv'])) + ')')
			par.paragraph_format.keep_with_next = True
		if 'export_773tg_str_mv' in record:
			par = doc.add_paragraph('In: ' + prep(record['export_773tg_str_mv'][0]))
			par.paragraph_format.keep_with_next = True
			for sub in record['export_773tg_str_mv'][1:]:
				par.paragraph_format.space_after = 5
				par = doc.add_paragraph(prep(sub))
				par.paragraph_format.keep_with_next = True
		if 'export_520a_str_mv' in record:
			par = doc.add_paragraph('[' + prep(' '.join(record['export_520a_str_mv'])) + ']')
			par.paragraph_format.keep_with_next = True
		if 'export_6xx_str_mv' in record:
			par = doc.add_paragraph()
			par.add_run('\u2022  ')
			par.add_run(prep('; '.join(record['export_6xx_str_mv'])))
			par.paragraph_format.keep_with_next = True
			par.paragraph_format.space_after = 5
		if 'export_787_str_mv' in record:
			for sub in record['export_787_str_mv']:
				par = doc.add_paragraph()
				par.add_run('\u279c  ')
				par.add_run(prep(sub))
				par.paragraph_format.keep_with_next = True
				par.paragraph_format.space_after = 5
		par = doc.add_paragraph()
	# foot
	#par = doc.add_paragraph(ADDRESS)
	#par.alignment = WD_ALIGN_PARAGRAPH.CENTER
	#foot = sections[-1].footer
	#foot.text = ADDRESS
	#write
	doc.save(ret)
	ret.seek(0)
	return ret

